"""Attachment upload service."""

from __future__ import annotations

import io
import json
import logging
import re
import uuid
from dataclasses import dataclass
from pathlib import Path

from fastapi import HTTPException, UploadFile, status

from app.core import settings

logger = logging.getLogger(__name__)

_ALLOWED_MIME_TO_TYPE = {
    "image/png": "image",
    "image/jpeg": "image",
    "image/webp": "image",
    "video/mp4": "video",
    "video/quicktime": "video",  # .mov
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "excel",  # .xlsx
    "application/vnd.ms-excel": "excel",  # .xls
    "application/vnd.ms-excel.sheet.macroenabled.12": "excel",  # .xlsm
    "application/vnd.ms-excel.sheet.binary.macroenabled.12": "excel",  # .xlsb
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",  # .docx
    "application/msword": "docx",  # .doc
    "text/plain": "txt",
    "text/csv": "excel",
    "application/csv": "excel",
}

_SAFE_NAME_RE = re.compile(r"[^a-zA-Z0-9._-]+")
_EXT_TO_TYPE = {
    ".png": "image",
    ".jpg": "image",
    ".jpeg": "image",
    ".webp": "image",
    ".mp4": "video",
    ".mov": "video",
    ".m4v": "video",
    ".xlsx": "excel",
    ".xls": "excel",
    ".xlsm": "excel",
    ".xlsb": "excel",
    ".csv": "excel",
    ".docx": "docx",
    ".doc": "docx",
    ".txt": "txt",
}


@dataclass
class SavedUpload:
    attachment_type: str
    attachment_url: str
    name: str
    mime_type: str
    size_bytes: int
    content: str | None = None  # For text-based files (txt, excel, docx)
    video_frames: list[str] | None = None


def _upload_dir_path() -> Path:
    """Resolve upload directory path consistently relative to backend root."""
    configured = Path(settings.UPLOAD_DIR)
    if configured.is_absolute():
        return configured
    backend_root = Path(__file__).resolve().parents[2]
    return (backend_root / configured).resolve()



def _sanitize_filename(filename: str) -> str:
    cleaned = _SAFE_NAME_RE.sub("_", filename).strip("._")
    return cleaned or "attachment"


def _attachment_type_for_mime(mime_type: str, filename: str | None = None) -> str:
    # Prefer extension when available because browser MIME for office files is often generic.
    ext = Path(filename or "").suffix.lower()
    if ext in _EXT_TO_TYPE:
        return _EXT_TO_TYPE[ext]

    normalized = mime_type.lower()
    attachment_type = _ALLOWED_MIME_TO_TYPE.get(normalized)
    if attachment_type:
        return attachment_type

    # Browsers sometimes send generic or unexpected MIME types for Office docs.
    if not attachment_type:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={
                "error": "unsupported_file_type",
                "message": f"Unsupported file type: mime='{normalized}', ext='{ext or 'unknown'}'.",
            },
        )
    return attachment_type


def _parse_txt_file(file_bytes: bytes) -> str:
    """Parse plain text file."""
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1", errors="replace")


def _parse_excel_file(file_bytes: bytes, filename: str) -> str:
    """Parse spreadsheet-like file and return JSON string."""
    ext = Path(filename).suffix.lower()
    try:
        import pandas as pd

        excel_file = io.BytesIO(file_bytes)
        if ext == ".csv":
            df = pd.read_csv(excel_file)
        elif ext in {".xlsx", ".xlsm"}:
            df = pd.read_excel(excel_file, engine="openpyxl")
        elif ext == ".xls":
            # Requires xlrd for legacy XLS parsing.
            df = pd.read_excel(excel_file, engine="xlrd")
        elif ext == ".xlsb":
            # Requires pyxlsb for binary Excel files.
            df = pd.read_excel(excel_file, engine="pyxlsb")
        else:
            # Last attempt for uncommon but compatible formats.
            df = pd.read_excel(excel_file)

        records = df.to_dict(orient="records")
        return json.dumps(records, indent=2, default=str)
    except Exception as e:  # noqa: BLE001
        logger.error("Failed to parse spreadsheet '%s': %s", filename, e)
        return json.dumps(
            {
                "error": "spreadsheet_parse_failed",
                "message": str(e),
                "filename": filename,
            }
        )


def _parse_docx_file(file_bytes: bytes) -> str:
    """Parse Word document (.docx) and return text content."""
    try:
        from docx import Document
        
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Also extract tables
        tables_data = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables_data.append(table_data)
        
        result = {
            "paragraphs": paragraphs,
            "tables": tables_data if tables_data else None,
        }
        return json.dumps(result, indent=2)
    except ImportError:
        logger.error("python-docx not installed")
        return json.dumps({"error": "python-docx library not installed"})
    except Exception as e:
        logger.error(f"Failed to parse DOCX: {e}")
        return json.dumps({"error": f"Failed to parse document: {str(e)}"})


def _extract_video_frames(file_bytes: bytes, num_frames: int = 3) -> list[str]:
    """Extract key frames from video and return as base64 data URIs."""
    try:
        import cv2
        import base64
        
        # Write to temp file (cv2 needs file path)
        temp_video = _upload_dir_path() / f"_temp_{uuid.uuid4().hex}.mp4"
        temp_video.write_bytes(file_bytes)
        
        try:
            cap = cv2.VideoCapture(str(temp_video))
            frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
            
            if frame_count == 0:
                return []
            
            frame_indices = [
                int(i * frame_count / (num_frames + 1)) 
                for i in range(1, num_frames + 1)
            ]
            
            frames_b64 = []
            for frame_idx in frame_indices:
                cap.set(cv2.CAP_PROP_POS_FRAMES, frame_idx)
                ret, frame = cap.read()
                
                if ret:
                    # Encode frame to JPEG bytes
                    ret, buffer = cv2.imencode(".jpg", frame)
                    if ret:
                        frame_bytes = buffer.tobytes()
                        b64 = base64.b64encode(frame_bytes).decode()
                        frames_b64.append(f"data:image/jpeg;base64,{b64}")
            
            cap.release()
            return frames_b64
        finally:
            temp_video.unlink(missing_ok=True)
            
    except ImportError:
        logger.warning("opencv-python (cv2) not installed, skipping frame extraction")
        return []
    except Exception as e:
        logger.error(f"Failed to extract video frames: {e}")
        return []


async def save_upload(file: UploadFile) -> SavedUpload:
    """Validate and persist an uploaded file to UPLOAD_DIR.
    
    For binary files (image, video): saves to disk.
    For text files (txt, excel, docx): parses content and stores with SavedUpload.
    For videos: also extracts key frames.
    """
    safe_name = _sanitize_filename(file.filename or "attachment")
    mime_type = (file.content_type or "application/octet-stream").lower()
    attachment_type = _attachment_type_for_mime(mime_type, safe_name)

    ext = Path(safe_name).suffix
    generated_name = f"{uuid.uuid4().hex}{ext}"

    upload_dir = _upload_dir_path()
    upload_dir.mkdir(parents=True, exist_ok=True)
    
    max_bytes = settings.MAX_UPLOAD_MB * 1024 * 1024
    size = 0
    file_bytes = io.BytesIO()

    # Read file into memory with size check
    while True:
        chunk = await file.read(1024 * 1024)
        if not chunk:
            break
        size += len(chunk)
        if size > max_bytes:
            raise HTTPException(
                status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                detail={
                    "error": "file_too_large",
                    "message": f"File exceeds {settings.MAX_UPLOAD_MB} MB limit.",
                },
            )
        file_bytes.write(chunk)

    await file.close()
    file_data = file_bytes.getvalue()
    
    # Parse content for text-based files
    content = None
    if attachment_type == "txt":
        content = _parse_txt_file(file_data)
    elif attachment_type == "excel":
        content = _parse_excel_file(file_data, safe_name)
    elif attachment_type == "docx":
        content = _parse_docx_file(file_data)
    
    # Save uploaded files to disk so /uploads can serve them.
    out_path = upload_dir / generated_name
    with out_path.open("wb") as f:
        f.write(file_data)

    # Extract video frames if video
    video_frames: list[str] | None = None
    if attachment_type == "video":
        video_frames = _extract_video_frames(file_data)

    return SavedUpload(
        attachment_type=attachment_type,
        attachment_url=f"/uploads/{generated_name}",
        name=safe_name,
        mime_type=mime_type,
        size_bytes=size,
        content=content,
        video_frames=video_frames,
    )
